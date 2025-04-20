import logging
from app.models.research import ResearchProject, ResearchOutline
from app.models.content import ResearchContent
import os
from flask import current_app
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from bidi.algorithm import get_display
import arabic_reshaper

logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting research papers"""

    def __init__(self):
        pass

    def generate_pdf(self, project_id, outline_id=None):
        """
        Generate a PDF with full Arabic support and larger fonts.
        """
        try:
            # Fetch project and outline
            project = ResearchProject.query.get(project_id)
            if not project:
                return {"error": "Project not found"}

            outline = (
                ResearchOutline.query.filter_by(
                    id=outline_id, project_id=project_id
                ).first()
                if outline_id
                else ResearchOutline.query.filter_by(
                    project_id=project_id, is_approved=True
                )
                .order_by(ResearchOutline.created_at.desc())
                .first()
            )
            if not outline:
                return {"error": "No approved outline found for this project"}

            # Check if the project is in Arabic
            is_arabic = project.language == "ar"

            # Configure ReportLab for Arabic
            if is_arabic:
                font_path = os.path.join(current_app.root_path, "static", "fonts")
                os.makedirs(font_path, exist_ok=True)

                arabic_font_path = os.path.join(
                    font_path, "NotoNaskhArabic-Regular.ttf"
                )
                arabic_bold_font_path = os.path.join(
                    font_path, "NotoNaskhArabic-Bold.ttf"
                )

                # Download Arabic fonts if missing
                if not os.path.exists(arabic_font_path) or not os.path.exists(
                    arabic_bold_font_path
                ):
                    import urllib.request

                    font_urls = {
                        "NotoNaskhArabic-Regular.ttf": "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoNaskhArabic/NotoNaskhArabic-Regular.ttf",
                        "NotoNaskhArabic-Bold.ttf": "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoNaskhArabic/NotoNaskhArabic-Bold.ttf",
                    }
                    for font_name, url in font_urls.items():
                        font_file_path = os.path.join(font_path, font_name)
                        if not os.path.exists(font_file_path):
                            urllib.request.urlretrieve(url, font_file_path)

                # Register Arabic fonts
                pdfmetrics.registerFont(TTFont("NotoNaskhArabic", arabic_font_path))
                pdfmetrics.registerFont(
                    TTFont("NotoNaskhArabicBold", arabic_bold_font_path)
                )

            # Helper to format Arabic text
            def format_text(text):
                if is_arabic and text:
                    reshaped_text = arabic_reshaper.reshape(text)
                    return get_display(reshaped_text)
                return text

            # Define styles with larger fonts
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                name="Title",
                parent=styles["Title"],
                fontName="NotoNaskhArabicBold" if is_arabic else "Helvetica-Bold",
                fontSize=30,
                alignment=1,
            )
            heading_style = ParagraphStyle(
                name="Heading",
                parent=styles["Heading2"],
                fontName="NotoNaskhArabicBold" if is_arabic else "Helvetica-Bold",
                fontSize=24,
                alignment=1,
            )
            normal_style = ParagraphStyle(
                name="Normal",
                parent=styles["Normal"],
                fontName="NotoNaskhArabic" if is_arabic else "Helvetica",
                fontSize=18,
                alignment=4,
                leading=24,
            )

            # Build PDF content
            elements = []
            elements.append(Paragraph(format_text(project.title), title_style))
            elements.append(Spacer(1, 0.5 * 72))  # 0.5 inch space

            # Table of Contents
            toc_title = (
                format_text("فهرس المحتويات") if is_arabic else "Table of Contents"
            )
            elements.append(Paragraph(toc_title, heading_style))
            elements.append(Spacer(1, 0.25 * 72))

            # Add sections to TOC
            outline_structure = outline.get_outline_structure()
            for section in outline_structure.get("sections", []):
                section_title = format_text(section.get("title", ""))
                elements.append(Paragraph(section_title, heading_style))
                elements.append(PageBreak())

            # Add content for each section
            for section in outline_structure.get("sections", []):
                section_title = format_text(section.get("title", ""))
                elements.append(Paragraph(section_title, heading_style))

                content_sections = ResearchContent.query.filter_by(
                    project_id=project_id, outline_id=outline.id
                ).all()
                section_content = next(
                    (
                        s.content
                        for s in content_sections
                        if s.section_title == section["title"]
                    ),
                    None,
                )
                if section_content:
                    elements.append(
                        Paragraph(format_text(section_content), normal_style)
                    )
                else:
                    empty_msg = (
                        format_text("لم يتم إنشاء محتوى لهذا القسم بعد.")
                        if is_arabic
                        else "Content not generated yet."
                    )
                    elements.append(Paragraph(empty_msg, normal_style))
                elements.append(PageBreak())

            # Save PDF
            export_dir = os.path.join(current_app.root_path, "static", "exports")
            os.makedirs(export_dir, exist_ok=True)
            filename = f"{project.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            filepath = os.path.join(export_dir, filename)

            doc = SimpleDocTemplate(
                filepath,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72,
            )
            doc.build(elements)

            return {"filepath": filepath, "filename": filename}

        except Exception as e:
            return {"error": str(e)}

    def generate_docx(self, project_id, outline_id=None):
        """
        Generate a DOCX file from project content

        Args:
            project_id: The ID of the research project
            outline_id: The ID of the specific outline to use (optional)

        Returns:
            dict: Information about the generated DOCX file
        """
        try:
            # Import docx library
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # Get the project
            project = ResearchProject.query.get(project_id)
            if not project:
                return {"error": "Project not found"}

            # Get the outline
            if outline_id:
                outline = ResearchOutline.query.filter_by(
                    id=outline_id, project_id=project_id
                ).first()
            else:
                # Get the latest approved outline
                outline = (
                    ResearchOutline.query.filter_by(
                        project_id=project_id, is_approved=True
                    )
                    .order_by(ResearchOutline.created_at.desc())
                    .first()
                )

            if not outline:
                return {"error": "No approved outline found for this project"}

            # Get all content sections
            content_sections = ResearchContent.query.filter_by(
                project_id=project_id, outline_id=outline.id
            ).all()

            # Create a dictionary to organize content by section title
            content_by_section = {}
            for section in content_sections:
                content_by_section[section.section_title] = {
                    "content": section.content,
                    "citations": section.get_citations(),
                }

            # Get the outline structure
            outline_structure = outline.get_outline_structure()

            # Create a new Document
            doc = Document()

            # Set RTL direction if language is Arabic
            is_rtl = project.language == "ar"

            # Add title
            title = doc.add_heading(project.title, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add Table of Contents heading
            if is_rtl:
                toc_title = "فهرس المحتويات"
            else:
                toc_title = "Table of Contents"

            doc.add_heading(toc_title, level=1)

            # Add TOC entries
            toc = doc.add_paragraph()

            # Add thesis statement to TOC if available
            if "thesis_statement" in outline_structure:
                if is_rtl:
                    toc.add_run("بيان الأطروحة").bold = True
                else:
                    toc.add_run("Thesis Statement").bold = True
                toc.add_run(" ........................... 1\n")

            # Add research questions to TOC if available
            if (
                "research_questions" in outline_structure
                and outline_structure["research_questions"]
            ):
                if is_rtl:
                    toc.add_run("أسئلة البحث").bold = True
                else:
                    toc.add_run("Research Questions").bold = True
                toc.add_run(" ........................... 2\n")

            # Add sections to TOC
            page_num = 3
            for section in outline_structure.get("sections", []):
                section_title = section.get("title", "")
                toc.add_run(f"{section_title}").bold = True
                toc.add_run(f" ........................... {page_num}\n")
                page_num += 1

            # Add references to TOC
            if is_rtl:
                toc.add_run("المراجع").bold = True
            else:
                toc.add_run("References").bold = True
            toc.add_run(f" ........................... {page_num}\n")

            # Add page break after TOC
            doc.add_page_break()

            # Add thesis statement if available
            if "thesis_statement" in outline_structure:
                if is_rtl:
                    doc.add_heading("بيان الأطروحة", level=1)
                else:
                    doc.add_heading("Thesis Statement", level=1)
                doc.add_paragraph(outline_structure["thesis_statement"])
                doc.add_page_break()

            # Add research questions if available
            if (
                "research_questions" in outline_structure
                and outline_structure["research_questions"]
            ):
                if is_rtl:
                    doc.add_heading("أسئلة البحث", level=1)
                else:
                    doc.add_heading("Research Questions", level=1)

                for i, question in enumerate(
                    outline_structure["research_questions"], 1
                ):
                    doc.add_paragraph(f"{i}. {question}")

                doc.add_page_break()

            # Add content for each section in the outline
            for section in outline_structure.get("sections", []):
                section_title = section.get("title", "")
                doc.add_heading(section_title, level=1)

                # Add the content if available
                if section_title in content_by_section:
                    html_content = content_by_section[section_title]["content"]

                    # Check if content is None or empty
                    if not html_content:
                        if is_rtl:
                            doc.add_paragraph("لم يتم إنشاء محتوى لهذا القسم بعد.")
                        else:
                            doc.add_paragraph(
                                "Content for this section has not been generated yet."
                            )
                    else:
                        # Parse HTML content
                        soup = BeautifulSoup(html_content, "html.parser")

                        # If no paragraphs found, add the whole content as one paragraph
                        if not soup.find_all(
                            ["p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "li"]
                        ):
                            doc.add_paragraph(html_content)
                        else:
                            # Process each paragraph
                            for p in soup.find_all(
                                [
                                    "p",
                                    "h1",
                                    "h2",
                                    "h3",
                                    "h4",
                                    "h5",
                                    "h6",
                                    "ul",
                                    "ol",
                                    "li",
                                ]
                            ):
                                if p.name.startswith("h"):
                                    doc.add_heading(p.get_text(), level=2)
                                elif p.name == "p":
                                    doc.add_paragraph(p.get_text())
                                elif p.name == "li":
                                    doc.add_paragraph(f"• {p.get_text()}")
                else:
                    if is_rtl:
                        doc.add_paragraph("لم يتم إنشاء محتوى لهذا القسم بعد.")
                    else:
                        doc.add_paragraph(
                            "Content for this section has not been generated yet."
                        )

                doc.add_page_break()

            # Add references section
            if is_rtl:
                doc.add_heading("المراجع", level=1)
            else:
                doc.add_heading("References", level=1)

            all_citations = []
            for section_data in content_by_section.values():
                all_citations.extend(section_data["citations"])

            # Remove duplicate citations
            unique_citations = {}
            for citation in all_citations:
                citation_id = citation.get("id", "")
                if citation_id and citation_id not in unique_citations:
                    unique_citations[citation_id] = citation

            # Add citations to DOCX
            for citation in unique_citations.values():
                doc.add_paragraph(citation.get("text", ""))

            # Create export directory if it doesn't exist
            export_dir = os.path.join(current_app.root_path, "static", "exports")
            os.makedirs(export_dir, exist_ok=True)

            # Create filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{project.id}_{timestamp}.docx"
            filepath = os.path.join(export_dir, filename)

            # Save the document
            doc.save(filepath)

            return {"filepath": filepath, "filename": filename}

        except ImportError:
            logger.error("python-docx library is not installed")
            return {
                "error": "python-docx library is not installed. Please install it with 'pip install python-docx'"
            }
        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            return {"error": str(e)}

    def paginate_content(
        self, sections_with_content, ordered_sections, max_words_per_page=250
    ):
        """
        Paginate section content to ensure each page has approximately max_words_per_page

        Args:
            sections_with_content: Dictionary of section content objects
            ordered_sections: List of section titles in order
            max_words_per_page: Maximum words per page (default: 250)

        Returns:
            List of page objects with paginated content
        """
        import re
        from bs4 import BeautifulSoup

        paginated_sections = []

        # Process each section
        for section_title in ordered_sections:
            if section_title not in sections_with_content:
                # Create an empty section for missing content
                paginated_sections.append(
                    {
                        "section_title": section_title,
                        "content": f'<p class="text-muted">Content not available</p>',
                        "word_count": 0,
                        "is_section_start": True,
                        "is_continuation": False,
                        "continues": False,
                    }
                )
                continue

            section = sections_with_content[section_title]
            content_html = section.content

            # Parse the HTML content
            soup = BeautifulSoup(content_html, "html.parser")

            # Get all text nodes and elements
            elements = []
            for element in soup.find_all(
                [
                    "p",
                    "h1",
                    "h2",
                    "h3",
                    "h4",
                    "h5",
                    "h6",
                    "ul",
                    "ol",
                    "blockquote",
                    "table",
                ]
            ):
                elements.append(element)

            if not elements:
                # Empty section
                paginated_sections.append(
                    {
                        "section_title": section_title,
                        "content": content_html,
                        "word_count": 0,
                        "is_section_start": True,
                        "is_continuation": False,
                        "continues": False,
                    }
                )
                continue

            # Start with section title on first page
            current_page_content = []
            current_page_word_count = 0
            is_section_start = True
            is_continuation = False

            for element in elements:
                # Count words in this element
                element_text = element.get_text()
                element_word_count = len(re.findall(r"\S+", element_text))

                # If adding this element would exceed the limit, create a new page
                if (
                    current_page_word_count + element_word_count > max_words_per_page
                    and current_page_content
                ):
                    # Create a page with current content
                    page_html = "".join(str(e) for e in current_page_content)
                    paginated_sections.append(
                        {
                            "section_title": section_title,
                            "content": page_html,
                            "word_count": current_page_word_count,
                            "is_section_start": is_section_start,
                            "is_continuation": is_continuation,
                            "continues": True,
                        }
                    )

                    # Reset for next page
                    current_page_content = []
                    current_page_word_count = 0
                    is_section_start = False
                    is_continuation = True

                # Add element to current page
                current_page_content.append(element)
                current_page_word_count += element_word_count

            # Add the last page for this section
            if current_page_content:
                page_html = "".join(str(e) for e in current_page_content)
                paginated_sections.append(
                    {
                        "section_title": section_title,
                        "content": page_html,
                        "word_count": current_page_word_count,
                        "is_section_start": is_section_start,
                        "is_continuation": is_continuation,
                        "continues": False,
                    }
                )

        return paginated_sections

    def generate_index_from_paginated(self, paginated_sections, language):
        """
        Generate table of contents with correct page numbers based on paginated sections
        """
        index = []
        current_page = 2  # Start from page 2 (after title page)
        section_pages = {}

        for section in paginated_sections:
            section_title = section["section_title"]

            # Only register the first occurrence of each section
            if section["is_section_start"] and section_title not in section_pages:
                section_pages[section_title] = current_page

                # Add to index
                index.append(
                    {
                        "title": section_title,
                        "page": current_page,
                        "indent": False,  # You may need to adjust this based on your outline structure
                    }
                )

            current_page += 1

        return index

    def save_pdf_annotations(self, pdf_path, annotations_data):
        """
        Save annotations to a PDF file

        Args:
            pdf_path: Path to the PDF file
            annotations_data: JSON data containing annotations

        Returns:
            str: Path to the annotated PDF file
        """
        try:
            import PyPDF2
            import json
            import os
            from datetime import datetime

            # Create a new filename for the annotated PDF
            filename = os.path.basename(pdf_path)
            base_name = os.path.splitext(filename)[0]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            annotated_filename = f"{base_name}_annotated_{timestamp}.pdf"

            # Use the static/exports directory
            export_dir = os.path.dirname(pdf_path)
            if "static" not in export_dir:
                export_dir = os.path.join(current_app.root_path, "static", "exports")

            annotated_pdf_path = os.path.join(export_dir, annotated_filename)

            # Read the original PDF
            reader = PyPDF2.PdfReader(pdf_path)
            writer = PyPDF2.PdfWriter()

            # Parse annotations data
            annotations = annotations_data
            if isinstance(annotations_data, str):
                annotations = json.loads(annotations_data)

            # Copy all pages from the original PDF
            for i in range(len(reader.pages)):
                page = reader.pages[i]
                writer.add_page(page)

                # Add annotations to the page if they exist
                page_key = f"page_{i+1}"
                if page_key in annotations:
                    page_annotations = annotations[page_key]

                    # Store annotations as PDF metadata
                    if "/Annots" not in page:
                        page["/Annots"] = []

                    # Create annotation dictionary
                    for annotation in page_annotations:
                        annot_dict = PyPDF2.generic.DictionaryObject()
                        annot_dict.update(
                            {
                                PyPDF2.generic.NameObject(
                                    "/Type"
                                ): PyPDF2.generic.NameObject("/Annot"),
                                PyPDF2.generic.NameObject(
                                    "/Subtype"
                                ): PyPDF2.generic.NameObject("/Text"),
                                PyPDF2.generic.NameObject(
                                    "/Contents"
                                ): PyPDF2.generic.createStringObject(
                                    json.dumps(annotation)
                                ),
                                PyPDF2.generic.NameObject(
                                    "/Rect"
                                ): PyPDF2.generic.ArrayObject(
                                    [
                                        PyPDF2.generic.FloatObject(0),
                                        PyPDF2.generic.FloatObject(0),
                                        PyPDF2.generic.FloatObject(0),
                                        PyPDF2.generic.FloatObject(0),
                                    ]
                                ),
                                PyPDF2.generic.NameObject(
                                    "/F"
                                ): PyPDF2.generic.NumberObject(28),
                            }
                        )

                        page["/Annots"].append(annot_dict)

            # Add document metadata with annotations
            metadata = PyPDF2.generic.DictionaryObject()
            metadata.update(
                {
                    PyPDF2.generic.NameObject(
                        "/Annotations"
                    ): PyPDF2.generic.createStringObject(json.dumps(annotations))
                }
            )
            writer.add_metadata(metadata)

            # Write the annotated PDF to file
            with open(annotated_pdf_path, "wb") as f:
                writer.write(f)

            return annotated_pdf_path

        except Exception as e:
            logger.error(f"Error saving PDF annotations: {str(e)}")
            # Return original PDF path if annotation fails
            return pdf_path

    def get_previous_exports(self, project_id):
        """
        Get a list of previously exported files for a project

        Args:
            project_id: The ID of the research project

        Returns:
            list: List of export file information
        """
        try:
            export_dir = os.path.join(current_app.root_path, "static", "exports")
            if not os.path.exists(export_dir):
                return []

            # Get all files that match the project ID pattern
            prefix = f"{project_id}_"
            exports = []

            for filename in os.listdir(export_dir):
                if filename.startswith(prefix) and filename.endswith(".pdf"):
                    filepath = os.path.join(export_dir, filename)
                    file_stats = os.stat(filepath)

                    # Extract timestamp from filename
                    timestamp_str = filename.replace(prefix, "").replace(".pdf", "")
                    try:
                        timestamp = datetime.strptime(timestamp_str, "%Y%m%d_%H%M%S")
                    except ValueError:
                        timestamp = datetime.fromtimestamp(file_stats.st_mtime)

                    exports.append(
                        {
                            "filename": filename,
                            "filepath": filepath,
                            "created_at": timestamp,
                            "size": file_stats.st_size,
                        }
                    )

            # Sort by creation time (newest first)
            exports.sort(key=lambda x: x["created_at"], reverse=True)
            return exports

        except Exception as e:
            logger.error(f"Error getting previous exports: {str(e)}")
            return []
